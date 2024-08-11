from django.shortcuts import render
import os
os.environ['KMP_DUPLICATE_LIB_OK'] = 'True'
import shutil
import fitz
from .models import *
from django.conf import settings
import cv2
from django.shortcuts import render
import layoutparser as lp
import glob
import pandas as pd
import numpy as np
from PIL import Image
from tqdm.auto import tqdm
from blend_modes import divide
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import easyocr
from PIL import Image, ImageFilter
import time
from azure.core.credentials import AzureKeyCredential
from azure.ai.formrecognizer import DocumentAnalysisClient
from dotenv import load_dotenv


load_dotenv()

# Set up Azure Form Recognizer client
key = os.getenv('key')
endpoint = os.getenv('endpoint')

document_analysis_client = DocumentAnalysisClient(endpoint=endpoint, credential=AzureKeyCredential(key))

# Create your views here.

def show_homepage(request):
    return render(request, 'index.html')

def process_image_for_table_extraction(image_path):
    with open(image_path, "rb") as f:
        image_data = f.read()

    poller = document_analysis_client.begin_analyze_document("prebuilt-layout", image_data)
    result = poller.result()

    tables_list = []
    for table in result.tables:
        data = [['' for _ in range(table.column_count)] for _ in range(table.row_count)]
        for cell in table.cells:
            data[cell.row_index][cell.column_index] = cell.content
        df = pd.DataFrame(data)
        df = df.transpose()
        tables_list.append(df)
    return tables_list

def render_table_as_html(tables_list):
    output_html_path = os.path.join(settings.MEDIA_ROOT, "tables.html")
    with open(output_html_path, 'w', encoding='utf-8') as f:
        f.write('<html><head><title>Extracted Tables</title></head><body>\n')
        for i, df in enumerate(tables_list):
            html_table = df.to_html(index=False)
            f.write(f'<h2>Table {i+1}</h2>\n')
            f.write(html_table)
            f.write('<br>\n')
        f.write('</body></html>')
    return output_html_path


def extract_and_display_tables(request):
    image_input = ImageInput.objects.first()  # Assuming you have a model instance with the image
    if not image_input:
        return render(request, 'index.html', {'error': 'No image found.'})
    
    image_path = image_input.image.path
    tables_list = process_image_for_table_extraction(image_path)
    html_path = render_table_as_html(tables_list)
    
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    return render(request, 'index.html', {'tables_html': html_content})

###################################################convert to doc starts here#########################################################

from docx import Document
import io
import cv2
from django.http import HttpResponse
from django.core.files.storage import default_storage

def get_y_coordinate(element):
    if hasattr(element, 'boundingRegions') and element.boundingRegions:
        return element.boundingRegions[0].boundingBox[1]
    elif hasattr(element, 'spans') and element.spans:
        return element.spans[0].offset
    return float('inf')

# def add_table_to_doc(table, doc):
#     table_doc = doc.add_table(rows=table.row_count, cols=table.column_count)
#     for cell in table.cells:
#         row = cell.row_index
#         col = cell.column_index
#         table_doc.cell(row, col).text = cell.content

def process_image(image_path):
    image = cv2.imread(image_path)
    doc = Document()

    with open(image_path, "rb") as f:
        image_data = f.read()

    poller = document_analysis_client.begin_analyze_document("prebuilt-layout", image_data)
    result = poller.result()

    page_elements = []
    for page in result.pages:
        if hasattr(page, 'lines'):
            for line in page.lines:
                page_elements.append(('line', get_y_coordinate(line), line))

    for paragraph in result.paragraphs:
        is_table_content = any(cell.content == paragraph.content for table in result.tables for cell in table.cells)
        if not is_table_content:
            page_elements.append(('paragraph', get_y_coordinate(paragraph), paragraph))

    for table in result.tables:
        page_elements.append(('table', get_y_coordinate(table), table))

    # Sort elements by their Y-coordinate to maintain order
    page_elements.sort(key=lambda x: x[1])


    for element_type, y_coord, element in page_elements:
        if element_type == 'paragraph':
            role = element.role
            content = element.content

            if role == "title":
                doc.add_heading(content, level=1)  # Adjust the level as needed
            elif role == "sectionHeading":
                doc.add_heading(content, level=2)  # Adjust the level as needed
            elif role == "pageHeader":
                doc.add_paragraph(content, style='Header')  # Add a header style if desired
            elif role == "pageFooter":
                doc.add_paragraph(content, style='Footer')  # Add a footer style if desired
            elif role == "footnote":
                doc.add_paragraph(content, style='Footnote')  # Add a footnote style if desired
            else:
                doc.add_paragraph(content)  # Default to a regular paragraph

        elif element_type == 'table':
            table_doc = doc.add_table(rows=element.row_count, cols=element.column_count)

            for cell in element.cells:
                row = cell.row_index
                col = cell.column_index
                table_doc.cell(row, col).text = cell.content

    return doc




def convert_image_to_docx(image_path):
    doc = process_image(image_path)

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def download_docx(request):
    if request.method == 'POST':
        if 'image' not in request.FILES:
            return HttpResponse("No image file uploaded.", status=400)

        image_file = request.FILES['image']
        image_path = default_storage.save(image_file.name, ContentFile(image_file.read()))
        image_path = os.path.join(settings.MEDIA_ROOT, image_path)

        docx_file = convert_image_to_docx(image_path)
        
        response = HttpResponse(docx_file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="output.docx"'
        return response
    
    return render(request, 'index.html', {'error': 'Invalid request method.'})






# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx import Document
# from docx.shared import Inches
# from scipy.stats import gaussian_kde
# from scipy.signal import argrelextrema
# import io

# def usePubLayNet(i, l:int):
#     publaynet = "lp://PubLayNet/ppyolov2_r50vd_dcn_365e/config"
#     pb_model= lp.PaddleDetectionLayoutModel(publaynet, device = 'gpu')
#     pb_layout = pb_model.detect(i)
#     pb_blocks = pb_layout._blocks

#     print(f'{len(pb_blocks)} pb_blocks: \n')
#     for item in pb_blocks:
#         print(item)

#     bboxes = []
#     for item in pb_blocks:
#         bbox = item.block
#         if item.type == 'Table' and l>0:
#             continue
#         else:
#             bboxes.append(((int(bbox.x_1),int(bbox.y_1),int(bbox.x_2),int(bbox.y_2)), item.type, item.score))
#     return bboxes

# def useTableLayout(i):
#     config = "lp://TableBank/ppyolov2_r50vd_dcn_365e/config"
#     tbmodel= lp.PaddleDetectionLayoutModel(config)
#     layout = tbmodel.detect(i)
#     blocks = layout._blocks
#     print("TABLE LAYOUT\n")
#     bboxes = []
#     for item in blocks:
#         bbox = item.block
#         bboxes.append(((int(bbox.x_1),int(bbox.y_1),int(bbox.x_2),int(bbox.y_2)), item.type, item.score))
#         print(item)
#     pass

#     return bboxes

# def find_columns(data:list):
  
#     data=np.array(data)
#     print(data)

#     # Estimate the density using kernel density estimation

#     try:
#         density = gaussian_kde(data)

#         # Generate x values for plotting
#         x = np.linspace(min(data), max(data), 1000)

#         # Calculate the density values for the x values
#         density_values = density(x)

#         # Find peaks in the density values
#         peaks_idx = argrelextrema(density_values, np.greater)[0]
#         peaks = x[peaks_idx]
#         print('peaks ', peaks)
#         return len(peaks)
#     except:
#         return 1

# def list_processing(list_string, doc:Document):
#     list_string = sorted(list_string, key=lambda x: (int(x[0][0][1]), int(x[0][0][0])))
#     i=0
#     word = list_string[0][1]
#     word_list = [(list_string[0][0][0][0],list_string[0][1])]

#     while i<len(list_string):
#         first_item = list_string[i]
#         first_y0=int(first_item[0][0][1])
#         word_height = int(first_item[0][2][1])-int(first_item[0][0][1])
#         word_height = word_height*.75
        
#         try:
#             second_item = list_string[i+1]
#             second_y0=int(second_item[0][0][1])
#         except:
#             second_item = (0,'')
#             second_y0 = first_y0+2*word_height
#             pass
        
#         if abs(first_y0-second_y0)<=word_height:
#             word_list.append((second_item[0][0][0],second_item[1]))
#         else:
#             word_list = sorted(word_list, key=lambda x:x[0])
#             word = ' '.join(item[1] for item in word_list)
#             doc.add_paragraph(word)        
#             try:
#                 word_list=[(second_item[0][0][0],second_item[1])]
#             except:
#                 pass
        
#         i+=1

# def get_image_description(image):
#     cropped_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
#     pil_image = Image.fromarray(cropped_rgb)
#     encoding = feature_extractor(image, return_tensors="pt")
#     keys=encoding.keys()

#     with torch.no_grad():
#         outputs = model(**encoding)
#     target_sizes = [pil_image.size[::-1]]
#     results = feature_extractor.post_process_object_detection(outputs, threshold=0.88, target_sizes=target_sizes)[0]
    
#     scores = results['scores']
#     labels = results['labels']
#     boxes = results['boxes']
    
#     converted_labels = [label_mapping[int(id)] for id in labels]
#     results['labels'] = converted_labels

#     return results

# def get_cell_coordinates_by_column(results):
#      # Extract labels and bounding boxes
#     labels = results['labels']
#     boxes = results['boxes']
    
#     # Initialize lists to store rows and columns
#     rows = []
#     columns = []
    
#     # Iterate over labels and boxes
#     for label, box in zip(labels, boxes):
#         if label == 'table row':
#             rows.append({'label': label, 'bbox': box})
#         elif label == 'table column':
#             columns.append({'label': label, 'bbox': box})
    
#     # Sort rows and columns by their Y and X coordinates, respectively
#     rows.sort(key=lambda x: x['bbox'][1])
#     columns.sort(key=lambda x: x['bbox'][0])
#     # Function to find cell coordinates
#     def find_cell_coordinates(row, column):
#         cell_box = [column['bbox'][0], row['bbox'][1], column['bbox'][2], row['bbox'][3]]
#         return cell_box
#     # Generate cell coordinates and count cells in each row
#     cell_coordinates = []
#     for column in columns:
#         column_cells = []
#         for row in rows:
#             cell_bbox = find_cell_coordinates(row, column)
#             column_cells.append({'row': row['bbox'], 'cell': cell_bbox})
#         # Sort cells in the row by Y coordinate
#         column_cells.sort(key=lambda x: x['row'][1])
#         # Append row information to cell_coordinates
#         cell_coordinates.append({'column': column['bbox'], 'cells': column_cells, 'cell_count': len(column_cells)})
#     # Sort columns from left to right
#     cell_coordinates.sort(key=lambda x: x['column'][0])
    
#     for item in cell_coordinates:
#         for cell_dict in item['cells']:
#             if isinstance(cell_dict['cell'], torch.Tensor):
#                 cell_dict['cell'] = [tensor.numpy() for tensor in cell_dict['cell']]
#                 cell_dict['cell'] = [arr.tolist() for arr in cell_dict['cell']]
#             else:
#                 cell_dict['cell'] = [tensor.numpy() for tensor in cell_dict['cell']]
#                 cell_dict['cell'] = [arr.tolist() for arr in cell_dict['cell']]
                
#     return cell_coordinates

# def apply_ocr(cell_coordinates, image):
#     # let's OCR row by row
#     ocr = PaddleOCR(use_angle_cls=True, lang='en')
#     data = dict()
#     max_num_rows = 0
#     cropped_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
#     pil_image = Image.fromarray(cropped_rgb)
    
#     for idx, column in enumerate(tqdm(cell_coordinates)):
#       column_text = []
#       for cell in column["cells"]:
#         cell_image = np.array(pil_image.crop(cell["cell"]))
#         # apply OCR
#         result = ocr.ocr(np.array(cell_image))
#         # print(result)
#         if result is None:
#             column_text.append(" ")
#         else:
#             if isinstance(result[0], list) and result[0]:
#                 text = " ".join([word[1][0] for word in result[0]])
#                 column_text.append(text)
#             else:
#                 column_text.append(" ")
#       if len(column_text) > max_num_rows:
#           max_num_rows = len(column_text)
#       data[idx] = column_text

#     # pad rows which don't have max_num_columns elements
#     # to make sure all rows have the same number of columns
#     for col, col_data in data.copy().items():
#         if len(col_data) != max_num_rows:
#           col_data = col_data + ["" for _ in range(max_num_rows - len(col_data))]
#         data[col] = col_data
#     return data

# def dataframe_to_docx(df, doc):
#     # Add table to document
#     table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
#     table.style = 'Table Grid'

#     # Add column names to the first row
#     for i, col_name in enumerate(df.columns):
#         table.cell(0, i).text = str(col_name)

#     # Add data to the table
#     for i, row in enumerate(df.iterrows()):
#         for j, value in enumerate(row[1]):
#             table.cell(i + 1, j).text = str(value)
            
# def page_to_document(file:FilePage):
#     # import easyocr
#     from easyocr import Reader
#     import gc
    
#     image_path = file.image.path
#     image = cv2.imread(image_path)
    
#     reader = Reader(['en', 'ne'])
#     doc = Document()
    
#     tlayout = useTableLayout(image)
#     playnet = usePubLayNet(image, len(tlayout))
#     bboxes= tlayout + playnet
#     x0_coords = [item[0][0] for item in bboxes]
#     no_of_cols = find_columns(x0_coords)
#     print(no_of_cols)
    
#     left_col=[]
#     right_col=[]
#     if no_of_cols==2:
#         mid_x0 = image.shape[1]/2
#         [left_col.append(item) if item[0][0]<mid_x0 else right_col.append(item) for item in bboxes]
#         left_col = sorted(left_col, key=lambda tup: tup[0][1])
#         right_col = sorted(right_col, key=lambda tup: tup[0][1])
        
#         cv2.line(
#             image,
#             (int(mid_x0), 0),
#             (int(mid_x0), int(image.shape[0])),
#             (0,0,0),
#             2
#         )
#     else:
#         left_col = sorted(bboxes, key=lambda tup:tup[0][1])
        
#     all_cols = left_col+right_col
#     words_by_para = []
    
#     for item in all_cols:
#         print(item)

#     for item in all_cols:
#         if item[2]<.75:
#             all_cols.remove(item)
            
#     for item in all_cols:
#         roi=image[item[0][1]:item[0][3], item[0][0]:item[0][2]] 

#         if item[1] == 'Figure':
#             if isinstance(roi, np.ndarray):
#         # Ensure roi is a valid image array
#                 if roi.shape[2] == 3:  # Check if it's RGB
#                     img = Image.fromarray(roi, mode='RGB')
#                 elif roi.shape[2] == 4:  # Check if it's RGBA
#                     img = Image.fromarray(roi, mode='RGBA')
#                 else:
#                     raise ValueError("Unsupported channel number in ROI array")
#                 img_io = BytesIO()
#                 img.save(img_io, format='PNG')
#                 img_io.seek(0) 
#                 doc.add_picture(img_io, width=Inches(2.0)) 
#             else:
#                 raise TypeError("roi must be a numpy array")

#         if item[1] in ['Text', 'Figure']:
#             easyocr_results = reader.readtext(roi)
#             # easyocr_results = []
#             print('\n', easyocr_results)
#             words = ' '.join(item[1] for item in easyocr_results)
#             p=doc.add_paragraph(words)
#             p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#             cv2.rectangle(
#                 image,
#                 (int(item[0][0]),int(item[0][1])),
#                 (int(item[0][2]), int(item[0][3])),
#                 (0,0,255),
#                 1
#             )
        
#         if item[1] == 'Title':
#             easyocr_results = reader.readtext(roi)
#             # easyocr_results = []
            
#             words = ' '.join(item[1] for item in easyocr_results)
#             doc.add_heading(words.title(), level=1)
#             cv2.rectangle(
#                 image,
#                 (int(item[0][0]),int(item[0][1])),
#                 (int(item[0][2]), int(item[0][3])),
#                 (0,0,255),
#                 1
#             )
        
#         if item[1] == 'List':
#             easyocr_results = reader.readtext(roi)
#             # easyocr_results = []

#             list_processing(easyocr_results, doc)
#             for list_item in easyocr_results:
#                 cv2.rectangle(
#                     image,
#                     (int(list_item[0][0][0])+ int(item[0][0]),int(list_item[0][0][1])+int(item[0][1])),
#                     (int(list_item[0][2][0])+int(item[0][0]),int(list_item[0][2][1])+int(item[0][1])),
#                     (0,255,0),
#                     1
#                 )
            
#         if item[1] == 'Table':
#             results = get_image_description(roi)
#             cell_coordinates = get_cell_coordinates_by_column(results)
#             data = apply_ocr(cell_coordinates, roi)
#             df = pd.DataFrame(data)

#             dataframe_to_docx(df, doc)
            
#             cv2.rectangle(
#                 image,
#                 (int(item[0][0]),int(item[0][1])),
#                 (int(item[0][2]), int(item[0][3])),
#                 (255,0,0),
#                 1
#             )
    
#     doc_buffer = io.BytesIO()
#     doc.save(doc_buffer)
#     doc_buffer.seek(0)
#     gc.collect()
    
#     return doc_buffer

# ##################################################CamScanner##############################################################
# #-------------------------------------------------IMAGE PROCESSING FOR PDF CONVERSION---------------------------------------------------
# def sharpen(image, strength=1):
#     # Define a sharpening kernel
#     kernel = np.array([[-1, -1, -1],
#                        [-1,  9, -1],
#                        [-1, -1, -1]])
    
#     # Create a kernel filter
#     kernel_filter = ImageFilter.Kernel((3, 3), kernel.flatten(), scale=strength)
    
#     # Apply filter to the image
#     sharpened_image = image.filter(kernel_filter)
    
#     return sharpened_image

# def resize_with_aspect_ratio(image_path, max_dimension=1500):
#     image = Image.open(image_path)
#     width, height = image.size
    
#     # if width < max_dimension or height<max_dimension:
#     #     return image
    
#     if width > height:
#         scaling_factor = max_dimension / width
#     else:
#         scaling_factor = max_dimension / height
    
#     new_width = int(width * scaling_factor)
#     new_height = int(height * scaling_factor)
#     resized_image = image.resize((new_width, new_height))
    
#     return resized_image

# def do_image_processing(x):
#     image = resize_with_aspect_ratio(x)
#     image_np = np.array(image)

#     yuv_image = cv2.cvtColor(image_np, cv2.COLOR_RGB2YUV)
#     y, u, v = cv2.split(yuv_image)
#     clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(5, 5))
#     y_eq = clahe.apply(y)
#     yuv_image_eq = cv2.merge((y_eq, u, v))
#     rgb_image_eq = cv2.cvtColor(yuv_image_eq, cv2.COLOR_YUV2BGR)
#     equalized_image = Image.fromarray(rgb_image_eq)

#     denoised_image_array = rgb_image_eq
#     denoised_image_array = cv2.bilateralFilter(denoised_image_array, d=-1, sigmaColor=50, sigmaSpace=50)
#     denoised_image_array =cv2.GaussianBlur(denoised_image_array, (3,3), 0)
#     denoised_image_array = cv2.bilateralFilter(denoised_image_array, d=-1, sigmaColor=25, sigmaSpace=25)


#     denoised_image = Image.fromarray(denoised_image_array)
#     sharpened_image = sharpen(denoised_image)

#     sharpened_image_array = np.array(sharpened_image)
#     threshold = (135,135,135)
#     mask = np.all(sharpened_image_array>threshold, axis=-1)
#     sharpened_image_array[mask] = (255,255,255)

#     threshold2 = (90,90,90)
#     mask = np.all(sharpened_image_array<threshold2, axis=-1)
#     sharpened_image_array[mask] = (0,0,0)

#     return sharpened_image_array.astype(np.uint8)
#     # thresholded_image = Image.fromarray(sharpened_image_array)
#     pass

# def de_shadow(image):
#     # splitting the image into channels
#     bA = image[:,:,0]
#     gA = image[:,:,1]
#     rA = image[:,:,2]

#     # dialting the image channels individually to spead the text to the background
#     dilated_image_bB = cv2.dilate(bA, np.ones((7,7), np.uint8))
#     dilated_image_gB = cv2.dilate(gA, np.ones((7,7), np.uint8))
#     dilated_image_rB = cv2.dilate(rA, np.ones((7,7), np.uint8))

#     # blurring the image to get the backround image
#     bB = cv2.medianBlur(dilated_image_bB, 21)
#     gB = cv2.medianBlur(dilated_image_gB, 21)
#     rB = cv2.medianBlur(dilated_image_rB, 21)

#     # blend_modes library works with 4 channels, the last channel being the alpha channel
#     # so we add one alpha channel to our image and the background image each
#     image = np.dstack((image, np.ones((image.shape[0], image.shape[1], 1))*255))
#     image = image.astype(float)
#     dilate = [bB,gB,rB]
#     dilate = cv2.merge(dilate)
#     dilate = np.dstack((dilate, np.ones((image.shape[0], image.shape[1], 1))*255))
#     dilate = dilate.astype(float)

#     # now we divide the image with the background image 
#     # without rescaling i.e scaling factor = 1.0
#     blend = divide(image,dilate,1.0)
#     blendb = blend[:,:,0]
#     blendg = blend[:,:,1]
#     blendr = blend[:,:,2]
#     blend_planes = [blendb,blendg,blendr]
#     blend = cv2.merge(blend_planes)
#     # blend = blend*0.85
#     blend = np.uint8(blend)

#     # returning the shadow-free image
#     return blend

# def camscanner_effect(image):
#     file_contents = image.file.read()
#     nparr = np.frombuffer(file_contents, np.uint8)
#     img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
#     # # image = cv2.resize(image, (1000,1000))
#     shadows_removed = de_shadow(img)          #de_shadow has better results
#     cv2.imwrite('test.png', shadows_removed)
    
#     gray = cv2.cvtColor(shadows_removed, cv2.COLOR_BGR2GRAY)
#     gray_thresh = cv2.threshold(gray, 0, 255,  cv2.THRESH_TOZERO|cv2.THRESH_OTSU )[1]
#     gray_thresh = cv2.threshold(gray_thresh, 225, 255 , cv2.THRESH_BINARY)[1]
    
#     return(gray_thresh)  

# def new_de_shadow(image):
#     # splitting the image into channels
#     bA = image[:,:,0]
#     gA = image[:,:,1]
#     rA = image[:,:,2]

#     # dialting the image channels individually to spead the text to the background
#     dilated_image_bB = cv2.dilate(bA, np.ones((7,7), np.uint8))
#     dilated_image_gB = cv2.dilate(gA, np.ones((7,7), np.uint8))
#     dilated_image_rB = cv2.dilate(rA, np.ones((7,7), np.uint8))

#     # blurring the image to get the backround image
#     bB = cv2.medianBlur(dilated_image_bB, 21)
#     gB = cv2.medianBlur(dilated_image_gB, 21)
#     rB = cv2.medianBlur(dilated_image_rB, 21)

#     # blend_modes library works with 4 channels, the last channel being the alpha channel
#     # so we add one alpha channel to our image and the background image each
#     image = np.dstack((image, np.ones((image.shape[0], image.shape[1], 1))*255))
#     image = image.astype(float)
#     dilate = [bB,gB,rB]
#     dilate = cv2.merge(dilate)
#     dilate = np.dstack((dilate, np.ones((image.shape[0], image.shape[1], 1))*255))
#     dilate = dilate.astype(float)

#     # now we divide the image with the background image 
#     # without rescaling i.e scaling factor = 1.0
#     blend = divide(image,dilate,.2)
#     blendb = blend[:,:,0]
#     blendg = blend[:,:,1]
#     blendr = blend[:,:,2]
#     blend_planes = [blendb,blendg,blendr]
#     blend = cv2.merge(blend_planes)
#     # blend = blend*0.85
#     blend = np.uint8(blend)

#     # returning the shadow-free image
#     return blend

# def new_process_image(image):
#     file_contents = image.file.read()
#     image_stream = io.BytesIO(file_contents)
#     y = do_image_processing(image_stream)
#     return y
    
# def process_image(image):
#     file_contents = image.file.read()
#     nparr = np.frombuffer(file_contents, np.uint8)
#     print(f'nparr: {nparr}')
#     img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
#     print(f'img: {img}')
#     # img = new_de_shadow(img)
#     img = de_shadow(img)
        
#     return img
    

# def images_to_pdf(image_paths):
#     output_pdf = BytesIO()

#     c = canvas.Canvas(output_pdf, pagesize=A4)
#     width, height = A4

#     for image_path in image_paths:
#         img = Image.open(image_path)
#         img_width, img_height = img.size

#         width_ratio = width / img_width
#         height_ratio = height / img_height
#         scale = min(width_ratio, height_ratio)

#         new_width = img_width * scale
#         new_height = img_height * scale

#         x_position = (width - new_width) / 2
#         y_position = (height - new_height) / 2

#         c.drawImage(image_path, x_position, y_position, new_width, new_height)
#         c.showPage()

#     c.save()

#     # Save the generated PDF file to a ContentFile
#     pdf_content = output_pdf.getvalue()
#     return pdf_content

# def images_to_a4_pdf(image_paths):
#     output_pdf = BytesIO()

#     c = canvas.Canvas(output_pdf, pagesize=A4)
#     page_width, page_height = A4

#     for image_path in image_paths:
#         img = Image.open(image_path)
#         img_width, img_height = img.size

#         # Stretch the image to fit the entire PDF page size
#         new_width = page_width
#         new_height = page_height

#         # Calculate position to start drawing the image (0,0 for full page)
#         x_position = 0
#         y_position = 0

#         # Draw the image stretched to fill the entire page
#         c.drawImage(image_path, x_position, y_position, new_width, new_height)
#         c.showPage()

#     c.save()

#     # Save the generated PDF file to a ContentFile
#     pdf_content = output_pdf.getvalue()
#     return pdf_content
    
# def create_pdf(images):
#     pdf_content = images_to_pdf(images)
#     return pdf_content
    